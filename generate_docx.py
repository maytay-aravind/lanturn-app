#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.dml.color import ColorFormat
import datetime

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.05

# ── Helpers ──
BRAND = RGBColor(0x1A, 0x1A, 0x1A)      # #1A1A1A
ACCENT = RGBColor(0xFF, 0xC1, 0x07)     # #FFC107 yellow
PINK = RGBColor(0xE9, 0x1E, 0x63)
GREY = RGBColor(0x64, 0x74, 0x8B)
LIGHT_BG = "F8F9FA"

def set_cell_bg(cell, color_hex):
    tblCell = cell._tc
    tblCellProperties = tblCell.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tblCellProperties.append(shd)

def add_para(text, bold=False, italic=False, size=9.5, color=None, align=None, space_after=4, bullet=False, level=0):
    p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
    if bullet and level>0:
        p.paragraph_format.left_indent = Inches(0.25*level)
    p.paragraph_format.space_after = Pt(space_after)
    if align: p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if italic: run.italic = True
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color: run.font.color.rgb = color
    return p

def add_heading_styled(text, level=1, color=BRAND, size=None, align=None, space_before=10, space_after=6, keep_with_next=True):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = keep_with_next
    if align: h.alignment = align
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.color.rgb = color
    run.bold = True
    if size: run.font.size = Pt(size)
    else:
        sizes = {1: 18, 2: 13, 3: 10.5, 4: 9.5}
        run.font.size = Pt(sizes.get(level, 10))
    # underline for H1/H2
    if level <=2:
        pPr = h._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'FFC107' if level==1 else 'E0E0E0')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return h

def add_table(headers, rows, col_widths=None, header_bg="1A1A1A", header_color="FFFFFF", font_size=7.5, zebra=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    # header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_bg(hdr_cells[i], header_bg)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(header_color)
        run.font.name = 'Calibri'
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, val in enumerate(row):
            cells[c_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if zebra and r_idx % 2 == 1:
                set_cell_bg(cells[c_idx], "F8F9FA")
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            # allow bold fragments via ** markers?
            if "**" in str(val):
                parts = str(val).split("**")
                for pi, part in enumerate(parts):
                    run = p.add_run(part)
                    run.bold = (pi % 2 == 1)
                    run.font.size = Pt(font_size)
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(0x33,0x33,0x33)
            else:
                run = p.add_run(str(val))
                run.font.size = Pt(font_size)
                run.font.name = 'Calibri'
                run.font.color.rgb = RGBColor(0x33,0x33,0x33)
            # center numeric cols
            if c_idx == 0 and len(headers)>3:
                pass
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table

def add_bullet_list(items, size=8.5, color=RGBColor(0x33,0x33,0x33), bold_prefix=False, space_after=1):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.left_indent = Inches(0.2)
        if bold_prefix and ":" in item:
            pre, post = item.split(":", 1)
            r1 = p.add_run(pre + ":")
            r1.bold = True
            r1.font.size = Pt(size)
            r1.font.name = 'Calibri'
            r1.font.color.rgb = color
            r2 = p.add_run(post)
            r2.font.size = Pt(size)
            r2.font.name = 'Calibri'
            r2.font.color.rgb = color
        elif "**" in item:
            parts = item.split("**")
            for pi, part in enumerate(parts):
                r = p.add_run(part)
                r.bold = (pi % 2 == 1)
                r.font.size = Pt(size)
                r.font.name = 'Calibri'
                r.font.color.rgb = color
        else:
            r = p.add_run(item)
            r.font.size = Pt(size)
            r.font.name = 'Calibri'
            r.font.color.rgb = color

def add_key_value(key, value, key_size=8.5, val_size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(key + "  ")
    r1.bold = True
    r1.font.size = Pt(key_size)
    r1.font.color.rgb = BRAND
    r1.font.name = 'Calibri'
    r2 = p.add_run(value)
    r2.font.size = Pt(val_size)
    r2.font.color.rgb = RGBColor(0x33,0x33,0x33)
    r2.font.name = 'Calibri'
    return p

# ────────────────────────────────────────────────────────────────
# TITLE PAGE
# ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Siva Sivani Degree College")
r.font.size = Pt(9)
r.font.color.rgb = GREY
r.font.name = 'Calibri'
r.bold = True
r.font.letter_spacing = Pt(1.2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
# Logo placeholder — big title
r = p.add_run("LanTURN")
r.font.size = Pt(42)
r.bold = True
r.font.color.rgb = BRAND
r.font.name = 'Calibri'
# yellow dot
r2 = p.add_run(" •")
r2.font.size = Pt(42)
r2.bold = True
r2.font.color.rgb = ACCENT
r2.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("AI-Powered Placement Platform")
r.font.size = Pt(13)
r.font.color.rgb = GREY
r.font.name = 'Calibri'
r.italic = True

# yellow line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("─" * 52)
r.font.size = Pt(8)
r.font.color.rgb = ACCENT
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Turning Learning into Career  •  Final Year Project Presentation")
r.font.size = Pt(9)
r.font.color.rgb = GREY
r.font.name = 'Calibri'
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Comprehensive Technical Guide  •  Stack  •  Features  •  Viva Q&A")
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(0x99,0x99,0x99)
r.font.name = 'Calibri'

# info box
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
table.columns[0].width = Inches(1.6)
table.columns[1].width = Inches(4.8)
cells = table.rows[0].cells
set_cell_bg(cells[0], "1A1A1A")
set_cell_bg(cells[1], "FFF8E1")
for c in cells:
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
p = cells[0].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("P\nR\nE\nS\nE\nN\nT\nA\nT\nI\nO\nN")
r.font.size = Pt(10)
r.bold = True
r.font.color.rgb = RGBColor.from_string("FFC107")
r.font.name = 'Calibri'
p = cells[1].paragraphs[0]
p.paragraph_format.left_indent = Inches(0.15)
lines = [
    ("Date  ", datetime.date.today().strftime("%B %d, %Y") + "  •  Academic Year 2025–26"),
    ("Team  ", "LanTURN Project Team  •  Siva Sivani Degree College"),
    ("Guide ", "Project Guide  •  Department of Computer Science"),
    ("Repo  ", "github.com/maytay-aravind/lanturn-app  •  Branch: master"),
    ("Live  ", "Frontend: Vercel  •  Backend: Render/Oracle  •  DB: Supabase  •  Auth: Firebase"),
]
for k,v in lines:
    p2 = cells[1].add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.space_before = Pt(1)
    r1 = p2.add_run(k + "  ")
    r1.bold = True
    r1.font.size = Pt(8)
    r1.font.color.rgb = BRAND
    r1.font.name = 'Calibri'
    r2 = p2.add_run(v)
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x33,0x33,0x33)
    r2.font.name = 'Calibri'
# remove first empty para in cell 1
if len(cells[1].paragraphs)>1:
    # first para is empty from table creation, remove its extra spacing
    pass
# clean top empty para
# Actually table's first para already set, second is our first added — we have an extra empty; ignore

doc.add_paragraph().paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("This document covers every layer of LanTURN — from tech choices to viva answers. Keep it open during your presentation.")
r.font.size = Pt(7.5)
r.italic = True
r.font.color.rgb = GREY
r.font.name = 'Calibri'

# ── Footer for title page ──
# Add page break after title
doc.add_page_break()

# ────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS (manual)
# ────────────────────────────────────────────────────────────────
add_heading_styled("Contents", level=1, size=16, space_before=0)
toc_items = [
    ("1", "Executive Summary — What is LanTURN?", "3"),
    ("2", "Technology Stack", "3"),
    ("3", "System Architecture & Project Structure", "6"),
    ("4", "Database Schema (Supabase PostgreSQL)", "7"),
    ("5", "Features — Complete Walkthrough", "9"),
    ("6", "API Specification", "14"),
    ("7", "Security, Rate-Limiting & Validation", "15"),
    ("8", "Deployment & DevOps", "16"),
    ("9", "Tools & Services Used", "16"),
    ("10", "Suggested Demo Flow (5–7 min)", "17"),
    ("11", "Viva — Questions & Model Answers", "17"),
    ("12", "Presentation Tips & Closing", "24"),
]
for num, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    # dotted leader via tab
    r1 = p.add_run(f"{num}.  {title}  ")
    r1.font.size = Pt(8.5)
    r1.font.name = 'Calibri'
    r1.font.color.rgb = BRAND
    # dots
    # use tab stop — simpler: add dots manually
    r2 = p.add_run("  " + "·" * 38 + f"  {pg}")
    r2.font.size = Pt(8)
    r2.font.color.rgb = GREY
    r2.font.name = 'Calibri'

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run("Tip:  Use Ctrl+F to jump to any section. Every heading is bookmarked.")
r.font.size = Pt(7)
r.italic = True
r.font.color.rgb = GREY
r.font.name = 'Calibri'

# ────────────────────────────────────────────────────────────────
# 1. EXECUTIVE SUMMARY
# ────────────────────────────────────────────────────────────────
add_heading_styled("1  Executive Summary — What is LanTURN?", level=1)

add_para("LanTURN = Turning Learning into Career — an AI-powered campus placement platform that connects students with the right employers based on skills, potential, and cultural fit. It digitizes the entire placement lifecycle: profile creation, job discovery, applications, employer screening, AI-assisted preparation, and admin governance.", size=8.5, color=RGBColor(0x33,0x33,0x33))

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Vision:  ")
r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = BRAND; r.font.name='Calibri'
r2 = p.add_run("Make campus placements intelligent, fair, and zero-cost — no fees, no premium walls, fully free-tier stack.")
r2.font.size = Pt(8.5); r2.font.color.rgb = RGBColor(0x33,0x33,0x33); r2.font.name='Calibri'

add_bullet_list([
    "**Problem solved:**  Manual resumes, scattered job posts, bias in shortlisting, no skill-gap feedback, opaque admin control.",
    "**Users:**  3 roles — Student, Employer (recruiter), Admin — each with a dedicated portal and guardrails.",
    "**AI-native:**  8+ AI tools (Gemini) for resume review, job match, skill gap, cover letter, interview prep, career chat, career DNA, company DNA, hiring assistant, skill arena.",
    "**Backend-first:**  Every feature is an API; web today, mobile/PWA tomorrow reuse 100%.",
    "**Free-tier only:**  Built to run on Firebase free + Supabase free + Gemini free + Vercel/Render free.",
], size=8, bold_prefix=False)

add_key_value("Tagline", "“Your Career Starts Here.”  — Landing page hero.")
add_key_value("Status", "v1 production-ready — deployed on Vercel (frontend) + Render (backend) + Supabase (DB/Storage) + Firebase (Auth).")

# Goals table
add_para("Goals & Non-Goals", bold=True, size=9, color=BRAND, space_after=2)
add_table(
    ["Type", "Item"],
    [
        ["Goal", "End-to-end placement lifecycle for students & employers"],
        ["Goal", "Reusable, well-documented REST API (backend-first)"],
        ["Goal", "AI differentiator — not just a CRUD job board"],
        ["Goal", "Free-tier deployment — cost ≈ $0"],
        ["Goal", "Modular, agent-friendly, RBAC-secure codebase"],
        ["Non-Goal", "Native mobile v1 (PWA planned next), payments, video interviews"],
    ],
    col_widths=[1.1, 5.3], font_size=7.5
)

# ────────────────────────────────────────────────────────────────
# 2. TECHNOLOGY STACK
# ────────────────────────────────────────────────────────────────
add_heading_styled("2  Technology Stack", level=1)

add_heading_styled("2.1  Frontend", level=2, size=11)
add_table(
    ["Layer", "Tech / Version", "Purpose"],
    [
        ["Framework", "React 18.3.1 (ESM)", "Component UI, hooks, code-splitting (lazy/Suspense)"],
        ["Build", "Vite 5.3.5 + @vitejs/plugin-react 4.3.1", "Dev server 5173, fast HMR, manualChunks vendor split"],
        ["Routing", "React Router 6.26.0", "SPA routing, guards (RequireAuth/Role/Onboarded), SPA rewrite"],
        ["State (server)", "TanStack Query 5.51.0", "Caching (stale 5m, gc 10m), prefetch, mutations"],
        ["State (global)", "AuthContext + LanguageContext", "Firebase user, role, i18n (EN/HI/TE)"],
        ["HTTP", "Axios 1.7.3 + apiClient interceptor", "Bearer ID token injection, error normalisation, unwrap()"],
        ["Auth", "Firebase Client SDK 10.12.4", "Google, Email/Pwd, Phone OTP, getIdToken()"],
        ["Styling", "Tailwind 3.4.7 + PostCSS 8.4.40", "brand/accent tokens, Nothing OS monochrome + brutalist shadows"],
        ["Icons", "Lucide 0.417", "No emojis — all icons are Lucide"],
        ["Animation", "Framer Motion 12.43", "Page transitions, Career Aisle timeline, modals"],
        ["Charts", "Recharts 3.10.1", "Donut, Pie, Bar, Area — dashboard & analytics"],
        ["Drag & Drop", "@dnd-kit 6.3/10.0/3.2", "Employer Kanban board for applicants"],
        ["Dates", "date-fns 3.6", "timeAgo(), deadline formatting"],
        ["Toasts", "react-hot-toast 2.4.1 (bottom-center)", "Success/error feedback"],
        ["Fonts", "Space Grotesk (headline) + Inter (body) + Space Mono", "Loaded via <link> in index.html"],
        ["Lint/Format", "ESLint 9.9", "npm run lint"],
    ],
    col_widths=[1.4, 2.1, 2.9], font_size=7
)

add_heading_styled("2.2  Backend", level=2, size=11)
add_table(
    ["Layer", "Tech / Version", "Purpose"],
    [
        ["Runtime", "Node ≥18, ESM (\"type\": \"module\")", "No CommonJS; native node --watch + node --test"],
        ["Framework", "Express 4.19.2", "REST API, trust proxy 1, 1 MB JSON limit"],
        ["Auth verify", "Firebase Admin 12.4.0", "verifyIdToken() only — no Firestore/Storage via Firebase"],
        ["Database", "Supabase JS 2.111 + PostgreSQL", "All tables via service_role (bypasses RLS)"],
        ["Validation", "Zod 3.23.8", "Strict schemas for every route (body/query/params)"],
        ["Security", "Helmet 7.1, CORS 2.8, Rate-Limit 7.4", "Headers, allowlist, 100/15s + AI 20/day"],
        ["Uploads", "Multer 2.2 (memory)", "PDF resume 5 MB parse; image 2 MB"],
        ["PDF", "pdf-parse 1.1.1", "Extract resumeText → cached in students.resume_text"],
        ["Logging", "Pino 9.3 + pino-pretty 13.1", "Structured JSON, requestId scoping"],
        ["HTTP log", "Morgan 1.10", "Dev request logs"],
        ["Env", "Custom .env loader + Zod (no dotenv)", "Zod-validated, fail-fast on bad env"],
        ["Imports", "Subpath map #config #services/* …", "Clean imports — no relative hell"],
        ["Tests", "node --test (native)", "npm test → tests/"],
    ],
    col_widths=[1.4, 2.1, 2.9], font_size=7
)

add_heading_styled("2.3  Database & Storage", level=2, size=11)
add_table(
    ["Component", "Tech", "Details"],
    [
        ["DB", "Supabase PostgreSQL", "12 tables + GIN indexes, triggers update_updated_at()"],
        ["ORM", "Supabase JS (no Prisma)", "Direct queries, rowToEntity mappers, snake→camel"],
        ["Storage", "Supabase Storage (3 buckets)", "resumes (private, 5 MB PDF), profile-pictures (public 2 MB), company-logos (public 2 MB)"],
        ["Auth store", "users table (uid PK = Firebase UID)", "Mirrors Firebase identity; role/student/employer/admin, profileComplete"],
        ["Legacy docs", "Firestore (docs-only)", "Actual runtime is PostgreSQL — docs are reference"],
    ],
    col_widths=[1.4, 2.1, 2.9], font_size=7.5
)

add_heading_styled("2.4  AI & External APIs", level=2, size=11)
add_table(
    ["Service", "Model / Key", "Used For"],
    [
        ["Gemini", "gemini-3.5-flash-lite → flash → 2.5 fallback; GEMINI_API_KEY (comma-separated, 429 retry)", "Resume review, job match, skill gap, cover letter, career DNA, company DNA, hiring assistant, skill test, resume extract"],
        ["Fallback", "DeepSeek (DEEPSEEK_API_KEY)", "Secondary LLM path"],
        ["Jobs", "Jooble API (JOOBLE_API_KEY)", "External job search proxy POST /jobs/external-search"],
        ["Quota", "ai_usage table + AI_RATE_LIMIT_PER_DAY=20", "Per-user daily AI cap; 429 when exceeded"],
    ],
    col_widths=[1.4, 2.2, 2.8], font_size=7.5
)

add_heading_styled("2.5  DevOps & Deployment", level=2, size=11)
add_table(
    ["Concern", "Choice", "Notes"],
    [
        ["Frontend host", "Vercel", "vercel.json rewrite /(.*) → /index.html (SPA)"],
        ["Backend host", "Render / Oracle (trust proxy 1)", "PORT=8080, health /api/health"],
        ["DB/Storage host", "Supabase Cloud", "PostgreSQL + Storage + service_role key (server-only)"],
        ["Auth host", "Firebase (prod) / Emulator (local 9099)", "FIREBASE_PROJECT_ID empty → emulator lanturn-dev"],
        ["Secrets", "Env only (never committed)", "Backend .env (gitignored) + Frontend VITE_ vars"],
        ["CI", "GitHub (master) + git push origin master", "AGENTS.md rule: always push immediately"],
        ["Monorepo", "backend/ + frontend/ + docs/ + free-claude-code (nested ignore)", "Independent npm installs"],
    ],
    col_widths=[1.4, 2.1, 2.9], font_size=7.5
)

add_heading_styled("2.6  All Libraries & Tools at a Glance", level=2, size=11)
add_para("For your slide — one consolidated list:", italic=True, size=7.5, color=GREY, space_after=2)
tools_rows = [
    ["Frontend libs", "React, Vite, React Router, TanStack Query, Axios, Firebase Client, Tailwind, Framer Motion, Recharts, DnD-Kit, Lucide, date-fns, react-hot-toast"],
    ["Backend libs", "Express, Firebase Admin, Supabase JS, Zod, Helmet, CORS, Rate-Limit, Morgan, Multer, pdf-parse, Pino"],
    ["Database", "PostgreSQL (Supabase), 12 tables, GIN indexes, JSONB, triggers"],
    ["Auth", "Firebase Auth (Google, Email/Pwd, Phone OTP) + Admin SDK verifyIdToken"],
    ["Storage", "Supabase Storage (3 buckets) + signed URLs (15 min upload, 1 hr download)"],
    ["AI", "Gemini (primary) + DeepSeek (fallback) + prompt engineering (temp 0.1–0.8, JSON mode)"],
    ["External", "Jooble Jobs API (external search)"],
    ["Dev tools", "VS Code, Git/GitHub, npm, Vite, ESLint, Postman/Thunder Client, Chrome DevTools"],
    ["Platforms", "Vercel (FE), Render/Oracle (BE), Supabase (DB), Firebase Console (Auth), GitHub"],
]
add_table(["Category", "Stack"], tools_rows, col_widths=[1.5, 4.9], font_size=7.5)

# ────────────────────────────────────────────────────────────────
# 3. ARCHITECTURE & STRUCTURE
# ────────────────────────────────────────────────────────────────
add_heading_styled("3  System Architecture & Project Structure", level=1)

add_heading_styled("High-Level Flow", level=2, size=11)
add_para("Browser (React + Vite + Firebase Client)  —Bearer ID token via apiClient→  Express (Helmet/CORS/RateLimit/Zod)  →  Firebase Admin verifyIdToken  →  Services → Repositories (Supabase PG) / Storage (signed URLs) / Gemini / Jooble. Responses are {data, meta:{requestId}}; errors are {error:{code,message,details}}.", size=8, color=RGBColor(0x33,0x33,0x33))

add_heading_styled("Backend Layering (strict)", level=3, size=9.5)
add_bullet_list([
    "**Routes**  →  thin wiring: path + middleware (authenticate, requireRole, validate, rateLimit) → controller.",
    "**Controllers**  →  ultra-thin: pull req.user/params/body/query, call one service, res.json({data}). No business logic.",
    "**Services**  →  all business rules, RBAC re-checks, transactions, notifications, AI prompts, error mapping.",
    "**Repositories / Clients**  →  I/O only: Supabase queries, Firebase Admin, Gemini/DeepSeek/Storage signed URLs, pdf-parse.",
    "**Middlewares**  →  auth (verify + upsert stub user + disabled check), rbac, validate (Zod), rateLimit, requestContext (requestId), error, notFound.",
    "**Utils**  →  ids (generateId 12-char base64url), pagination (cursor base64), logger (Pino), AppError (typed HTTP codes).",
], size=7.5)

add_heading_styled("Frontend Layering", level=3, size=9.5)
add_bullet_list([
    "**App.jsx**  →  lazy routes + Suspense (DeepamLoader) + guards (RequireAuth → RequireOnboarded → RequireRole).",
    "**Contexts**  →  AuthContext (firebaseUser, role, isOnboarded) + LanguageContext (EN/HI/TE).",
    "**Services**  →  apiClient (Axios + Bearer) + typed modules (student/employer/job/application/ai/...), unwrap() reads data.",
    "**React Query**  →  server state: staleTime 5m, gcTime 10m, prefetch in Layout.jsx, invalidate on mutation.",
    "**Components**  →  ui/* (Modal, Skeleton, SkillsInput…), ai/* (CareerDNAPanel, RadarChart…), layout (Navbar, DeepamLoader).",
    "**No business logic in components** — every rule lives on the backend.",
], size=7.5)

add_heading_styled("Monorepo Layout", level=2, size=11)
# use a text block instead of table for layout
layout_lines = [
    "root/  AGENTS.md  start-dev.bat  cors.json",
    "├─ backend/  src/server.js (entry)  src/app.js (Express wiring)",
    "│           src/config/ (env Zod, constants)  src/firebase/ (Admin)  src/supabase/ (client + migrations 001-008 + setup-storage)",
    "│           src/middlewares/  src/utils/  src/schemas/ (Zod)  src/repositories/  src/clients/ (gemini, storage, jooble)",
    "│           src/services/  src/controllers/  src/routes/  src/data/careerRoadmaps.js  scripts/seed.js",
    "├─ frontend/  src/main.jsx  src/App.jsx  vite.config.js  tailwind.config.js  vercel.json",
    "│            src/contexts/  src/lib/ (apiClient, utils)  src/services/  src/hooks/  src/components/  src/pages/ (student/employer/admin/public)",
    "├─ docs/  15 markdown files (overview, architecture, API spec, schema, roadmap…)",
    "└─ free-claude-code/  (nested git repo — ignored for LanTURN work)",
]
for ln in layout_lines:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(ln)
    r.font.size = Pt(7)
    r.font.name = 'Consolas'
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)

add_heading_styled("Ports & Envs", level=3, size=9.5)
add_table(
    ["Service", "Port / Var", "Default / Notes"],
    [
        ["Backend API", "PORT → 8080", "Health http://localhost:8080/api/health"],
        ["Frontend Vite", "5173", "npm run dev in frontend/"],
        ["Firebase Auth Emu", "9099", "FIREBASE_AUTH_EMULATOR_HOST when PROJECT_ID empty"],
        ["Firestore Emu", "8081 (AGENTS) / 8080 (docs)", "Legacy — runtime uses Supabase PG"],
        ["Storage Emu", "9199", "Legacy — Supabase Storage is primary"],
        ["Backend env", "backend/.env (gitignored)", "Copy from .env.example; secret keys here"],
        ["Frontend env", "frontend/.env (gitignored)", "VITE_API_URL + VITE_FIREBASE_* + VITE_USE_EMULATOR"],
    ],
    col_widths=[1.4, 1.6, 3.4], font_size=7
)

# ────────────────────────────────────────────────────────────────
# 4. DATABASE SCHEMA
# ────────────────────────────────────────────────────────────────
add_heading_styled("4  Database Schema — Supabase PostgreSQL (Migrations 001–008)", level=1)
add_para("All tables live in Supabase PG. Docs/DATABASE_SCHEMA.md is legacy Firestore reference — real schema is migrations/001–008. Every table has updated_at trigger via update_updated_at(). Service role bypasses RLS.", size=7.5, italic=True, color=GREY)

add_table(
    ["Table", "PK / Key Columns", "Purpose"],
    [
        ["users", "uid (Firebase UID)", "Identity mirror; email, role (student/employer/admin), profileComplete, status"],
        ["students", "uid FK users", "1:1 profile — personal/academic/professional/social JSONB, searchable_skills GIN, resume_text/keywords, certificates"],
        ["employers", "uid FK users", "1:1 company — company_name, industry, location/hr_contact JSONB, logo, size/ceo/foundedYear/benefits/technologies, company_dna"],
        ["jobs", "job_id", "Postings — title/desc, required_skills GIN, job_type, experience, salary JSONB, work_mode, status (draft/active/verified/paused/closed/removed), verified_by_admin"],
        ["applications", "application_id = student_job", "Student×Job unique; snapshots (name/photo/resume/skills), status + history, employer_id denorm"],
        ["notifications", "notification_id", "In-app + email; user_id, type, title/body/link, read, email_status, channel"],
        ["analytics_events", "id UUID", "Admin counts: job_posted, application_submitted, etc. by type+created_at"],
        ["platform_config", "id='default' singleton", "signup_enabled, maintenance_mode, ai_daily_limit"],
        ["chat_threads", "thread_id", "AI chats — user_id, mode (career_guidance/employer_hiring/general), last_message preview"],
        ["chat_messages", "id UUID", "Thread messages — thread_id, role user/assistant, content"],
        ["career_roadmaps", "roadmap_id", "Enrollments — student_id+domain_id unique, domain_title"],
        ["roadmap_progress", "id UUID", "Per-topic completion — roadmap_id+stage_index+topic_index unique"],
        ["candidate_matches", "id UUID, (job_id,student_id) unique", "AI match 0–100 — scores, missingSkills, recommendations, reason"],
    ],
    col_widths=[1.5, 1.9, 3.0], font_size=7
)

add_heading_styled("Key Relationships", level=3, size=9.5)
add_bullet_list([
    "users 1—1 students, 1—1 employers, 1—* jobs (employer_id), 1—* applications (student_id), 1—* notifications, 1—* chat_threads 1—* chat_messages, 1—* career_roadmaps 1—* roadmap_progress, jobs 1—* applications, jobs 1—* candidate_matches *—1 users (student).",
    "GIN indexes: students.searchable_skills, jobs.required_skills; composite: jobs(status,created), applications(student_id,created)/(job_id,created)/(employer_id), notifications(user_id,created)/(user_id,read) partial.",
    "Triggers: BEFORE UPDATE on users/students/employers/jobs/applications/chat_threads → update_updated_at().",
], size=7)

add_heading_styled("Storage Buckets (Supabase Storage)", level=3, size=9.5)
add_table(
    ["Bucket", "Public", "Limit / MIME", "Path Example", "Kind"],
    [
        ["resumes", "No (private)", "5 MB / PDF only", "resumes/{uid}/resume-{uuid}.pdf", "resume + certificate"],
        ["profile-pictures", "Yes", "2 MB / PNG/JPEG/WEBP", "photos/{uid}/profilePhoto-{uuid}.jpg", "profilePhoto"],
        ["company-logos", "Yes", "2 MB / PNG/JPEG/WEBP", "logos/{uid}/companyLogo-{uuid}.png", "companyLogo"],
    ],
    col_widths=[1.2, 0.7, 1.3, 2.0, 1.2], font_size=7
)
add_para("Signed upload URLs (createSignedUploadUrl, 15 min) + signed download (1 hr) for resumes/certificates; public URLs for images. Files never proxied through Node — only URLs stored in DB.", size=7, color=GREY, italic=True)

# ────────────────────────────────────────────────────────────────
# 5. FEATURES
# ────────────────────────────────────────────────────────────────
add_heading_styled("5  Features — Complete Walkthrough", level=1)
add_para("Every feature below maps to routes → controllers → services → repositories. Guards: RequireAuth (any), RequireOnboarded, RequireRole(student/employer/admin), requireProfileComplete.", size=7.5, color=GREY, italic=True)

add_heading_styled("5.1  Public & Auth", level=2, size=11)
add_heading_styled("Landing Page  (/)", level=3, size=9.5)
add_bullet_list([
    "Fixed nav (LanTURN + dot accent), How It Works / Features / Stats anchors, Sign In vs Return to Dashboard.",
    "Hero  Your Career Starts Here. + subtitle + Get Started Free → /login?mode=signup + Watch Demo + dashboard mock preview.",
    "Trusted strip (brand logos), Stats (500+ / 120+ / 95%), How It Works (3 staggered cards: Create Profile, AI Matches, Get Placed), Features (Resume AI, Smart Matching, Career Assistant, Skill Analysis), CTA dark banner (email → Join Now), Footer (Privacy/Terms/Contact).",
], size=7.5)
add_heading_styled("Login  (/login,  /login?mode=signup)", level=3, size=9.5)
add_bullet_list([
    "Email/Password (register+login, 6-char min, show/hide eye, mapFirebaseError) + Forgot password (reset email).",
    "Google OAuth (select_account) — signin requires existing role else logout; signup → /onboarding.",
    "Phone OTP (RecaptchaVerifier + requestPhoneOtp → 6-digit verify with Change number).",
    "Split UI: left charcoal Nothing OS illustration, right card with GoogleLogo, OR divider, recaptcha-container.",
], size=7.5)
add_heading_styled("Onboarding  (/onboarding) — 3-step wizard", level=3, size=9.5)
add_bullet_list([
    "Step 1 Language (flag cards EN/HI/TE → LanguageContext).  Step 2 Role (Student GraduationCap vs Employer Briefcase).",
    "Step 3 Profile: Student → personal (name/phone/city), academic (college/degree/branch/year/cgpa/city), professional.skills (comma), social (linkedin/github/portfolio). Employer → companyName/industry/size/website/city/description.",
    "Actions: Skip (empty profile → still creates row) vs Complete Setup → POST /auth/onboarding {role,profile} → refreshSession() → role dashboard. All fields optional.",
], size=7.5)
add_heading_styled("Admin Login  (/admin-login)", level=3, size=9.5)
add_para("Separate page — authenticate then POST /auth/admin-login checks email == ADMIN_EMAIL → promotes to admin; else error. Redirects to /admin.", size=7.5)

add_heading_styled("5.2  Student Portal  (/dashboard, /jobs, /applications, /profile, /ai, /career-aisle, /notifications)", level=2, size=11)

add_heading_styled("Dashboard  (/dashboard) — Insights", level=3, size=9.5)
add_bullet_list([
    "Header banner (Welcome back {name}) + Find Jobs → /jobs; prefetch (Layout) for fast first paint.",
    "3 stat cards: Applications count + View Applications; Profile Strength donut (25% each: name/college/resume/skills) + Complete/NeedsInfo badge; Top Skills chips (5 + +N).",
    "Application Pipeline — Recharts Pie (52/78) segments submitted grey / reviewing slate / shortlisted black / hired emerald / rejected red / withdrawn muted, legend + total center.",
    "Recent Applications list (5, avatar initial, timeAgo, status badges) → View All.",
    "Career Aisle Progress (if enrolled): grid of SVGDonuts per roadmap (percentComplete, green if 100 else black) → navigate to /career-aisle with roadmapId.",
    "Top Hiring Partners — toggle Distribution (horizontal dual-pill tracks jobWidth/appWidth vs maxCombined, rank #1-5, logo, verified check) vs Bar Overview (BarChart jobs #1A1A1A vs apps #B0B0B0, tooltip). Filter by metric.",
], size=7)

add_heading_styled("Jobs — Unified  (/jobs,  /job-search → redirect)", level=3, size=9.5)
add_bullet_list([
    "Tabs inside frosted card: Internal Jobs (Briefcase) vs External Jobs (Globe, Jooble).",
    "Internal: Search q + jobType filter + limit 50 + cursor pagination Load More; sorted by computeJobMatch score (skillMatchRatio*75 + bonuses for resume/projects/experience/>5 skills, or text-match fallback). Job cards: logo/initial, title/company, verified ShieldCheck, matchScore% badge (Highly Qualified ≥80 / Good Fit ≥60 / Partial ≥40 / Low), meta (location/type/salary/deadline), description clamp, skills pills, View Details or Applied/Expired. Detail dialog (portal, backdrop-blur): header (logo, badges, Verified), toggle Job Details vs Company Profile View, Qualification Match card (score + matched/missing chips), Quick Stats (location/salary/openings/deadline), About/Responsibilities/Requirements/ Skills/Education/Benefits/Stipend, Company Profile (logo, HQ/size/CEO/foundedYear/website/linkedin/email, CompanyDNAPanel, technologies/benefits/culture), footer Posted • applicants + Apply Now → applicationService.apply().",
    "External: keywords + location + page, joobleService.search, Prev/Next, popular pills (Software Engineer etc.), cards with Building2 avatar, title→link, snippet HTML, Apply CTA → external.",
], size=7)

add_heading_styled("Applications  (/applications) — Kanban + List", level=3, size=9.5)
add_bullet_list([
    "Header toggle Board (LayoutGrid) vs List (List); summary colored pills per status.",
    "Kanban (read-only): 5 columns Applied (Clock blue) → Reviewing (Eye amber) → Shortlisted (UserCheck emerald) → Hired (CheckCircle2 violet) → Rejected (XCircle red), vertical ArrowDown chain, cards grid 1-4 cols (avatar, jobTitle, company, timeAgo, coverLetter, Withdraw if submitted/pending).",
    "List: card per app (avatar, jobTitle, company, applied timeAgo, coverLetter, status badge, Withdraw). Mutations: applicationService.listMine(50), withdraw (DELETE) with confirm → invalidate.",
], size=7)

add_heading_styled("Profile  (/profile) — Read + Edit", level=3, size=9.5)
add_bullet_list([
    "Header My Profile; banner gradient + Edit Profile toggle (Pencil/X) + avatar (photo or initial) + Camera → fileRef image ≤2MB → uploadService.uploadFile('profilePhoto').",
    "View mode: name, degree·branch·college, badges (role, Class of year, CGPA, GitHub #0A66C2, LinkedIn slate), detail rows (phone/city/college/skills/GPA), skills pills + SkillMedalBadge (gold/silver/bronze/basic), links pills, resume inline (status, View signedUrl, AI Score → /ai?tab=review, Auto-fill → aiService.extractResume merges into state, Update/Upload PDF ≤5MB, progress bar), certificates grid (View signed cert, Rename, Remove, upload PDF/JPG/PNG/WEBP ≤5MB).",
    "Edit mode (Sections): Personal (phone/city/state), Academic (college/degree/branch/year/CGPA), Professional (SkillsInput Enter/comma), Social (GitHub/LinkedIn/Portfolio/HR), Resume dropzone, Certificates. Sticky Save → studentService.updateMe() with Zod toast.",
], size=7)

add_heading_styled("AI Career Assistant  (/ai?tab=chat|dna|review|match|arena)", level=3, size=9.5)
add_bullet_list([
    "Header Sparkles + 5 tabs synced to ?tab:",
    "① Career Chat — h-520 flex, Bot/User bubbles, markdown (bold/headers/lists), 4 suggestion chips, threadId persistence, aiService.careerChat({message,threadId,mode:career_guidance}) <150w bullets.",
    "② Career DNA — Generate Career DNA → aiService.careerDna() then getCareerDna cache; OverallScoreRing (96px, emerald/amber/red), careerField + candidateLevel, RadarChart 340px (6 dims, clickable), Strengths/Weaknesses/Recommendations grids, Re-analyze, per-dimension modal (reason, suggestions).",
    "③ Resume Review — targetRole input (auto-predicted), Analyze → aiService.reviewResume({targetRole}) → ScoreRing (score + ATS), keywords badges, strengths (emerald), weaknesses+fixes (amber+Chevron), missing keywords (red), predictedRole autofill.",
    "④ Job Match — select job (jobService.list 50), Match → aiService.matchJob({jobId}) → ScoreRing (local computeJobMatch or matchScore), experienceFit badge (strong/partial/weak), summary, matched/missing skills.",
    "⑤ Skill Arena — Trophy: pick skill (searchable_skills + resumeKeywords), rating slider, Generate 3 open-ended Qs (tiered by rating: 90-100 Expert architecture, 80-89 Advanced…), answer, Evaluate → aiService skill-test/evaluate → medal gold/silver/bronze/basic/none, X/3, feedback, persisted skillMedals.",
], size=7)

add_heading_styled("Career Aisle  (/career-aisle) — 57 Roadmaps", level=3, size=9.5)
add_bullet_list([
    "Catalog via roadmapService: 57 domains (frontend, backend, aiml, devops…) grouped All/Software/Mobile/AI/Cloud/Product/Management/Emerging; icons per domain.",
    "DomainPickerModal (framer backdrop-blur): catalog search (Flutter/Quantum/AWS) + difficulty filter, 2-col cards (icon, category, desc, months, stages) → Preview (gradient hero, Enroll Zap, stage-by-stage curriculum with topics/resources/projects) or Select (enroll). Header Scan Resume with AI → ResumeGapAnalyzerModal.",
    "Timeline per enrolled roadmap: stats (percentComplete, completed/total, stages, months, Remove), center alternating timeline (vertical line + progress height by percentComplete, left/right StagePanels 380px with gradient accent, timeframe, topic rows (checkbox line-through when done → toggleTopic), project gradient card, resources). Node pill (Stage X / ✓ Complete, yellow vs emerald), end marker Journey Complete!.",
    "Resume Analyzer: upload PDF (5 MB) → POST /roadmaps/analyze-resume → resumeAnalyzer (Gemini) returns matchScore, extractedSkills, matchedTopicKeys, missingKeywords, recommendedStageIndex, summary → Sync to auto-complete matched topics + jump.",
], size=7)

add_heading_styled("Notifications  (/notifications)", level=3, size=9.5)
add_bullet_list([
    "notificationService.list(50) + unreadCount, Mark All Read (CheckCheck), NotifCard (ring if unread, icon per type Briefcase/FileText/AlertCircle/Sparkles/Bell, body, timeAgo, Check → markRead). Polling/fallback when onSnapshot not used.",
], size=7)

add_heading_styled("5.3  Employer Portal  (/employer/*)", level=2, size=11)

add_heading_styled("Dashboard  (/employer/dashboard)", level=3, size=9.5)
add_bullet_list([
    "Banner (Welcome {companyName}, industry) + Post Job; 4 stat cards (Active Jobs, Total Applicants, Shortlisted, Hired conversionRate) with Briefcase/Users/UserCheck/Award.",
    "Charts: Donut 160px Recruitment Pipeline (submitted #6B6B6B / reviewed #B0B0B0 / shortlisted #333 / accepted #1A1A1A / rejected #D62828) + legend+total; MiniArea cubic bezier 30-day applications (area gradient #1A1A1A 12%→0% + dots + tooltip).",
    "Metrics: Hiring Metrics bars (shortlisted/hired/rejected vs total) + conversion card; Job Overview grid (Active/Paused/Closed/Total).",
    "AI Recommended Candidates (4) — CandidateMatchCard with medals 🥇🥈🥉 + matchScore/jobTitle; Company DNA preview (generate/regenerate) via CompanyDNAPanel; Recent Applications list (photo/initial, name, jobTitle, status badge, timeAgo).",
], size=7)

add_heading_styled("Jobs  (/employer/jobs) — CRUD + AI", level=3, size=9.5)
add_bullet_list([
    "Header My Jobs + Post New Job / Cancel; form (card, slide-up) with SkillsInput pills: title*/role/department/jobType*/workMode/expLevel; description* textarea + Generate with AI (Sparkles → employerService.aiGenerateJobDesc → fills desc/responsibilities/requirements/skills/niceToHave); responsibilities/requirements (one per line), requiredSkills/education, location (city/state/country, remote), compensation (min/max + currency INR/USD/EUR/GBP + negotiable / stipend for internship), openings, deadline, benefits, status (active/draft).",
    "List: skeleton/empty, cards with StatusBadge (active green/paused yellow/closed/default/draft blue) + verified ShieldCheck + workMode badges, meta (city/department/experience, salary, skills, timeAgo/deadline), actions: {applicationCount} Applicants → /employer/jobs/:jobId/applicants, Pause/Play toggle, Edit (fills form), Delete confirm. Mutations invalidate jobs/mine + analytics.",
], size=7)

add_heading_styled("Applicants  (/employer/jobs/:jobId/applicants) — Pipeline Control", level=3, size=9.5)
add_bullet_list([
    "Header Back to Jobs + title + count; toggle Board vs List; status filter pills (All + per-status if count>0).",
    "ProfileModal: studentService.getPublic(studentId) → avatar, name/headline/location, social (GitHub/LinkedIn/Portfolio), education (college/degree/branch/year/cgpa), skills pills, projects (link+tech badges), experience, certifications.",
    "Actions per applicant: statusMutation → applicationService.updateStatus({status: reviewed/shortlisted/accepted/rejected}) + toast + invalidate; handles reviewed (Eye amber), shortlisted (UserCheck emerald), accepted (Award brand), rejected (UserX red); Resume → getResumeUrl signed 1 hr.",
    "Kanban columns (Applied blue / Reviewing amber / Shortlisted emerald / Hired violet / Rejected red) with DnD onDragEnd → statusMutation; cards (avatar, name, timeAgo, matchScore%, skillsSnapshot, View Profile/Resume).",
    "List cards: avatar, name+status badge, coverLetter, skillsSnapshot, timeAgo, Eye/View Profile + FileText/Resume, status buttons, CandidateMatchCard below if match exists (matchesMap from employerService.getJobMatches).",
], size=7)

add_heading_styled("AI Hiring Assistant  (/employer/ai-assistant) — Threaded Chat", level=3, size=9.5)
add_bullet_list([
    "Suggested prompts (Users Find React interns, Briefcase full-stack, Zap deployed projects, GraduationCap 2026). ThreadSidebar: New Conversation + search + list (MessageSquare, title, last preview, Clock, active ring, delete confirm). Main: Bot gradient header + Online pulse, messages area (welcome Sparkles + 2-col suggestions vs history max-w-3xl, Bot/User avatars, markdown, TypingIndicator 3-dot bounce), auto-scroll, auto-resize textarea (48→128px), Send/Loader. Backend employerAI.service builds context from active jobs + candidatePool (skills/projects/exp/edu) + candidateMatches then calls Gemini temp 0.5.",
], size=7)

add_heading_styled("Profile  (/employer/profile) — Company", level=3, size=9.5)
add_bullet_list([
    "View: cover gradient + cubes.png, 32px logo (or Building2) + Camera (≤2MB image → uploadFile companyLogo), verified emerald pill vs amber pending, Grid2 Company Details (name/industry/size/employees/website) vs Leadership & Contact (CEO/foundedYear/HQ/email/phone/HR), Culture & Achievements (About, culture, technologies Cpu pills, benefits Heart emerald, achievements Star, LinkedIn #0A66C2).",
    "Edit: same sections as editable inputs + TagsInput (pill + X, Enter/Plus to add) for technologies/benefits/achievements, LinkedIn URL, Company DNA panel (generate/regenerate → employerService.generateCompanyDna, 24h cooldown). Sticky Save bar → employerService.updateMe().",
], size=7)

add_heading_styled("Notifications  (/employer/notifications)", level=3, size=9.5)
add_para("Same pattern as student: type icons (application_received/application_status/job_applicant/system), markRead/markAllRead.", size=7.5)

add_heading_styled("5.4  Admin Portal  (/admin) — Shield", level=2, size=11)
add_bullet_list([
    "Auth: /admin-login email must equal ADMIN_EMAIL → promote to admin. Portal header Shield gradient + Admin Portal subtitle; tabs Dashboard / Post Verification / User Management (hash routing).",
    "Dashboard Tab: 8 stat cards (Total Users, Students indigo, Employers purple, Active Jobs emerald, Applications cyan, Verified Jobs blue, Hired amber, Pending orange) + Recent Posts (5, logo, status badge) + Recent Users (8, photo/initial, email, role badge, counts). Data adminService.getAnalyticsSummary + listJobs + listUsers.",
    "Post Verification Tab: summary 3 cards (Total/Verified/Pending), Search + status filter, cards (logo, title+status+Verified badge, company, type•Posted•applicants) + Verify (ShieldCheck) if pending else Verified pill + Change Status select → PATCH /admin/jobs/:jobId/verify or /status. JobsTab variant (hidden tab) same as table grid.",
    "User Management Tab: header {n} Users + Search + role/status filters; grid 1/2/3 cards (photo/initial, name/email, role badge, Active/Disabled pill, Joined timeAgo) + Actions dropdown (Change Role student/employer/admin disabled if current; Disable/Enable). Mutations PATCH /admin/users/:uid/status|role with click-outside handling.",
    "Also: Platform Config GET/PATCH /platform/config (signup_enabled, maintenance_mode, ai_daily_limit) and Analytics series by type/days.",
], size=7)

add_heading_styled("5.5  Cross-Cutting Features", level=2, size=11)
add_bullet_list([
    "**Career Aisle (57 roadmaps):**  Static CAREER_DOMAINS in backend/src/data/careerRoadmaps.js; per-domain color/gradient/icon/estimatedMonths/stages→topics→resources→projects; enroll/remove/toggleTopic + resume-analyzer gap sync.",
    "**i18n:**  EN/HI/TE via LanguageContext; Language switcher in Navbar/Onboarding.",
    "**Recharts:**  Student pipeline Donut + Employer miniArea + Admin could use but currently stat cards.",
    "**DnD:**  Employer applicants Kanban uses @dnd-kit.",
    "**DeepamLoader:**  Custom South Indian lamp loader for Suspense fallbacks.",
    "**EmptyState / Skeleton:**  Consistent skeletons + EmptyState with icon+title+description+CTA across apps/jobs/notifications.",
    "**Toasts:**  react-hot-toast bottom-center (error/success).",
    "**Top Companies:**  GET /employers/top-companies (public) ranked by totalJobs+totalApplications via JS aggregation.",
], size=7.5)

# ────────────────────────────────────────────────────────────────
# 6. API SPEC
# ────────────────────────────────────────────────────────────────
add_heading_styled("6  API Specification", level=1)
add_para("Base https://<host>/api   Auth Bearer <Firebase ID token>   Success {data, meta:{requestId}}   Error {error:{code,message,details}}   Pagination ?limit&cursor → {items,nextCursor}   RateLimit 429.", size=7, color=GREY, italic=True)

api_groups = [
    ("Auth", "GET /auth/session (me)  •  POST /auth/onboarding {role,profile}  •  POST /auth/logout  •  POST /auth/admin-login"),
    ("Students", "GET /students/me  •  PATCH|PUT /students/me  •  GET /students/me/resume-url  •  GET /students/me/certificate-url?url=  •  GET /students/:uid (public, employer)"),
    ("Employers", "GET /employers/me  •  PATCH|PUT /employers/me  •  GET /employers/me/analytics  •  GET /employers/:uid (public)  •  GET /employers/top-companies?limit  •  GET /employers/jobs/:jobId/matches  •  POST /employers/ai/*"),
    ("Jobs", "GET /jobs?q&jobType&limit&cursor  •  GET /jobs/me/all (employer)  •  GET /jobs/:jobId  •  POST /jobs (employer)  •  PATCH /jobs/:jobId  •  DELETE /jobs/:jobId  •  Status: draft/active/verified/paused/closed/removed"),
    ("Applications", "POST /jobs/:jobId/applications {coverLetter}  •  GET /applications (student mine)  •  GET /jobs/:jobId/applications (employer)  •  GET /applications/:id  •  PATCH /applications/:id/status {status}  •  DELETE /applications/:id (withdraw)  •  GET /applications/:id/resume-url"),
    ("Notifications", "GET /notifications?limit&cursor  •  GET /notifications/unread-count  •  PATCH /notifications/:id/read  •  POST /notifications/read-all"),
    ("AI (student)", "POST /ai/resume-extract  •  /ai/resume-review {targetRole?}  •  /ai/resume-match {jobId}  •  /ai/skill-gap {jobId}  •  /ai/skill-test/generate {skill,rating}  •  /ai/skill-test/evaluate  •  GET /ai/skill-medals  •  /ai/cover-letter {jobId,tone}  •  /ai/career-chat {message,threadId,mode}  •  /ai/career-dna (POST+GET)  •  GET /ai/threads + /ai/threads/:id/messages"),
    ("Roadmaps", "GET /roadmaps/domains  •  GET /roadmaps/domains/:id  •  GET /roadmaps/me  •  POST /roadmaps/me/enroll {domainId}  •  DELETE /roadmaps/me/:roadmapId  •  PATCH /roadmaps/me/:roadmapId/progress  •  POST /roadmaps/analyze-resume (multer PDF)  •  POST /roadmaps/sync-resume-progress"),
    ("Uploads", "POST /uploads/sign {fileName,contentType,size,purpose→ bucket}  •  POST /uploads/commit {key,url} → updates DB + extracts pdf keywords"),
    ("External", "POST /jobs/external-search {keywords,location,page} (Jooble proxy, hides API key)"),
    ("Admin", "GET /admin/users?role&status&limit  •  PATCH /admin/users/:uid/status {status}  •  PATCH /admin/users/:uid/role {role}  •  GET /admin/jobs  •  PATCH /admin/jobs/:jobId/status  •  PATCH /admin/jobs/:jobId/verify  •  GET /admin/analytics/summary  •  GET /admin/analytics/series?type&days  •  GET|PATCH /platform/config"),
    ("Health", "GET /health, GET /health/ready, GET /version"),
]
for title, endpoints in api_groups:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(title + "  ")
    r1.bold = True; r1.font.size = Pt(7.5); r1.font.color.rgb = BRAND; r1.font.name='Calibri'
    r2 = p.add_run(endpoints)
    r2.font.size = Pt(7); r2.font.color.rgb = RGBColor(0x33,0x33,0x33); r2.font.name='Calibri'

add_para("Error codes: VALIDATION_ERROR 400, UNAUTHENTICATED 401, FORBIDDEN 403, NOT_FOUND 404, CONFLICT 409 (duplicate application), UNPROCESSABLE 422 (no resume), RATE_LIMITED 429, INTERNAL 500, UPSTREAM_ERROR 502. All validated by Zod; all routes return requestId for tracing.", size=7, color=GREY)

# ────────────────────────────────────────────────────────────────
# 7. SECURITY
# ────────────────────────────────────────────────────────────────
add_heading_styled("7  Security, Rate-Limiting & Validation", level=1)
add_table(
    ["Concern", "How LanTURN Handles It"],
    [
        ["Auth", "Firebase ID JWT (~1h TTL), verifyIdToken on every request, disabled-user check, photo/name sync, upsert stub on first login"],
        ["RBAC", "requireRole(...roles) + requireProfileComplete; services re-check ownership (employer can only manage own jobs/applications)"],
        ["Validation", "Zod strict schemas on every body/query/params; strips unknown keys; AppError with field details"],
        ["Rate limits", "General 100/15s per IP, auth stricter burst, AI 20/day (ai_usage) + in-memory window"],
        ["Files", "Resume PDF ≤5MB, images ≤2MB, mime allowlist, per-user path {uid}/, signed URLs (15m upload, 1h download), never proxy"],
        ["Transport", "Helmet headers, CORS allowlist (5173), Bearer not cookies, secrets only in env (gitignored)"],
        ["Prompt injection", "System-first prompts, delimiters, trailing “ignore instructions” defense, blocked tokens"],
        ["IDOR / Mass assign", "Ownership checks in services; USER_LOCKED_FIELDS [role,status,email,uid] blocked; student can’t set employer fields"],
        ["Privacy", "Resumes visible only to owner + applied employer (signed URL check); public student/employer views are sanitized"],
        ["Logging", "Pino + requestId per request; no PII/secrets in logs; central error envelope"],
    ],
    col_widths=[1.4, 5.0], font_size=7
)

# ────────────────────────────────────────────────────────────────
# 8. DEPLOYMENT & DEVOPS
# ────────────────────────────────────────────────────────────────
add_heading_styled("8  Deployment & DevOps", level=1)
add_table(
    ["Step", "Command / Config", "Notes"],
    [
        ["Env setup", "cp backend/.env.example backend/.env  &&  cp frontend/.env.example frontend/.env", "Set VITE_API_URL=http://localhost:8080/api, VITE_USE_EMULATOR, GEMINI keys, SUPABASE_URL+SERVICE_ROLE"],
        ["Emulators", "firebase emulators:start", "Auth 9099, Firestore 8081, Storage 9199 (Supabase is primary)"],
        ["Backend dev", "cd backend && npm install && node src/supabase/setup-storage.js && npm run seed && npm run dev", "node --watch on 8080; setup-storage creates 3 buckets; seed (legacy, may need rerun)"],
        ["Frontend dev", "cd frontend && npm install && npm run dev", "Vite 5173; win: start-dev.bat"],
        ["Build", "frontend: npm run build → dist/ (manualChunks); backend: node src/server.js", "Vite chunk warns at 300kb"],
        ["Deploy FE", "Vercel (vercel.json SPA rewrite)", "Auto-deploy on git push master"],
        ["Deploy BE", "Render / Oracle VM (trust proxy 1)", "Set env vars in dashboard; health /api/health"],
        ["DB migrate", "Run backend/src/supabase/migrations/*.sql via Supabase SQL Editor", "001→008 in order; setup-storage once"],
        ["Seed data", "Run seed.js (currently Firestore-era — rewrite for Supabase if needed)", "Contains 2 employers + 3 students + 3 jobs for demo"],
    ],
    col_widths=[1.3, 2.7, 2.4], font_size=7
)

# ────────────────────────────────────────────────────────────────
# 9. TOOLS
# ────────────────────────────────────────────────────────────────
add_heading_styled("9  Tools & Services Used", level=1)

add_heading_styled("Development Tools", level=3, size=9.5)
add_table(
    ["Tool", "Version / Notes", "Used For"],
    [
        ["VS Code / Cursor", "Latest", "Primary IDE"],
        ["Git + GitHub", "master branch, push immediately (AGENTS.md)", "Version control; repo maytay-aravind/lanturn-app"],
        ["Node.js", "≥18 (ESM)", "Runtime for both BE and FE tooling"],
        ["npm", "10+", "Install + scripts (dev/build/lint/test)"],
        ["Vite", "5.3.5", "Frontend bundler/HMR"],
        ["ESLint", "9.9", "Linting (npm run lint)"],
        ["Postman / Thunder Client", "—", "API testing (Bearer token via Firebase print token)"],
        ["Chrome DevTools", "—", "Frontend debug, network, React Query"],
        ["Firebase CLI / Emulators", "firebase-tools 15.24", "Local Auth emulator + token mint"],
        ["Supabase Dashboard", "Cloud", "SQL Editor (migrations), Storage buckets, Table view"],
        ["Vercel Dashboard", "—", "FE deploys + env vars + SPA rewrite"],
        ["Render / Oracle Console", "—", "BE deploys + env vars + logs"],
        ["python-docx", "1.2", "This doc generation (bonus)"],
    ],
    col_widths=[1.5, 1.7, 3.2], font_size=7
)

add_heading_styled("Platforms & Services", level=3, size=9.5)
add_table(
    ["Service", "Role", "Cost"],
    [
        ["Firebase Auth", "Identity (Google/Email/Phone) + Admin SDK verify", "Free tier"],
        ["Supabase", "PostgreSQL + Storage (3 buckets) + JS client", "Free tier"],
        ["Gemini API", "All 8 AI features (primary LLM)", "Free tier (comma-separated keys)"],
        ["DeepSeek", "Fallback LLM", "Free tier"],
        ["Jooble API", "External job search", "Free / API key"],
        ["Vercel", "Frontend hosting + SPA rewrite", "Free tier"],
        ["Render / Oracle", "Backend hosting", "Free tier"],
        ["GitHub", "Code + AGENTS.md + pushes", "Free"],
    ],
    col_widths=[1.5, 2.7, 2.2], font_size=7
)

# ────────────────────────────────────────────────────────────────
# 10. DEMO FLOW
# ────────────────────────────────────────────────────────────────
add_heading_styled("10  Suggested Demo Flow (5–7 minutes)", level=1)
add_para("Start at / → /login → onboarding → pick one full user journey; narrate while clicking.", italic=True, size=7.5, color=GREY)
demo_steps = [
    "**0:00 Landing** — Scroll hero → Trusted → Stats → How It Works → Features → CTA; show fixed nav, parallax dots, DeepamLoader on lazy load.",
    "**0:45 Auth** — Show Login split screen: toggle Sign In/Sign Up, Email vs Phone, Google button, Forgot password link, recaptcha. Quick Google login.",
    "**1:15 Onboarding** — Step through Language → Role (Student vs Employer cards) → Profile form (Student personal/academic/skills/social vs Employer company); Skip vs Complete.",
    "**1:45 Student: Profile** — Show avatar upload (Camera ≤2MB), resume upload (PDF ≤5MB + Auto-fill extract), skill medals, certificates, Edit toggle + Save bar.",
    "**2:15 Student: Jobs** — Internal tab: search “React” + filter jobType, show matchScore badge (Highly Qualified 80+), card → dialog (Qualification Match + Company Profile toggle + Apply). External tab: Jooble search, external Apply link. Show deadline Expired greyscale.",
    "**2:50 Student: Applications** — Toggle Board vs List, show kanban columns + Withdraw on pending, timeAgo.",
    "**3:10 Student: AI Assistant** — Carousel 5 tabs: Chat (send “roadmap for MERN” → bulleted 150w answer), Career DNA (Generate → radar + score), Resume Review (targetRole → score + ATS + strengths/fixes), Job Match (pick job → matched/missing), Skill Arena (pick skill → Generate 3 Qs → Evaluate → medal).",
    "**3:50 Student: Career Aisle** — Show 57 Domains modal (All/Software/AI… + search + difficulty), Preview stage curriculum, Enroll, then Timeline (alternating panels, topic checkboxes, project cards, progress %). Click Resume Analyzer → PDF → gap summary.",
    "**4:20 Student: Dashboard** — Banner, 3 stat cards, Pie pipeline, Recent apps, Career Aisle donut grid, Top Hiring Partners (Distribution dual-pill vs Bar).",
    "**4:40 Employer: Dashboard** — 4 stats, Donut + MiniArea (30-day), Hiring Metrics bars, AI Recommended Candidates (medals), Recent apps + photos.",
    "**5:00 Employer: Jobs** — Post New Job + Generate with AI (Sparkles fills desc/requirements/skills), list cards with Pause/Delete/Edit, Applicants count → Applicants page.",
    "**5:20 Employer: Applicants** — Filter pills, Avatar + photo, ProfileModal (education/skills/projects/exp), status buttons (Review→Shortlist→Hired/Rejected), Handle Resume signed URL, Kanban drag or List + match card.",
    "**5:40 Employer: Hiring Assistant** — Show threads sidebar, suggested prompts, send “find top React intern” → context-aware reply with candidateNumbers, typing indicator.",
    "**5:55 Employer: Profile** — CompanyDetails + Leadership grid, Culture/Tech/Benefits/Achievements pills, logo Camera upload, Generate Company DNA + regeneration.",
    "**6:10 Admin (/admin-login → /admin)** — Dashboard 8 stats + Recent Posts/Users, Post Verification (Verify + Change Status), Users (role/status filters + Actions dropdown Disable/Enable + Change Role).",
    "**6:30 Close** — Mention free-tier architecture, PWA next, Q&A.",
]
for s in demo_steps:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    parts = s.split("**")
    for pi, part in enumerate(parts):
        r = p.add_run(part)
        r.bold = (pi % 2 == 1)
        r.font.size = Pt(7)
        r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(0x33,0x33,0x33)

# ────────────────────────────────────────────────────────────────
# 11. VIVA Q&A
# ────────────────────────────────────────────────────────────────
add_heading_styled("11  Viva — Questions & Model Answers", level=1)
add_para("~40 questions grouped by topic. Answers are concise, viva-style — expand with an example if asked to elaborate. Every answer is grounded in your actual code.", size=7.5, color=GREY, italic=True)

qa_groups = [
    ("A. Stack & Basics (7 Qs)", [
        ("What is the tech stack of LanTURN?",
         "Monorepo: frontend React 18 + Vite 5 + Tailwind + React Router 6 + TanStack Query 5 + Axios + Firebase Client SDK 10 + Framer Motion + Recharts + DnD-Kit + Lucide + date-fns; backend Node ≥18 ESM + Express 4 + Firebase Admin 12 + Supabase JS 2 (PostgreSQL) + Zod 3 + Helmet/CORS/RateLimit + Multer + pdf-parse + Pino; DB Supabase PostgreSQL (12 tables), Storage 3 buckets, AI Gemini (primary) + DeepSeek + Jooble API; deploys Vercel (FE) + Render/Oracle (BE)."),
        ("Why Vite instead of CRA/Next.js?",
         "Vite gives sub-second HMR, native ESM, tiny config, manualChunks for vendor splitting and long-term caching — perfect for an SPA that doesn’t need SSR. Next would add SSR complexity we don’t need v1; CRA is deprecated."),
        ("Why React 18 and not Angular/Vue?",
         "Ecosystem + hiring pool, fine-grained control over auth flow, easiest interop with Firebase JS SDK, and TanStack Query gives us server-state caching without Redux boilerplate."),
        ("Why Express over NestJS/Fastify?",
         "Express is minimal, well-known, free-tier friendly, and our layering (routes→controllers→services→repos) is explicit without Nest decorators. Fastify perf gain is negligible vs clarity for a college project."),
        ("What does \"type\": \"module\" mean?",
         "Backend and frontend both use native ESM — import/export, top-level await capable, no require(). Enables subpath imports (#services/*) and keeps bundling clean."),
        ("What is TanStack Query doing for you?",
         "It owns all server state: caching (stale 5m, gc 10m), deduping, retries (1), background refetch off, prefetch in Layout, and cache invalidation after every mutation (e.g., apply → invalidate ['applications','mine'] + analytics). No Redux needed."),
        ("What is Axios vs fetch?",
         "We use Axios via apiClient — interceptors auto-inject Firebase Bearer token, normalise error envelope to Error.code/details, and unwrap() reads data. Fetch would need manual boilerplate."),
    ]),
    ("B. Architecture & Patterns (6 Qs)", [
        ("Explain the backend layering.",
         "Strict: routes (path+middleware) → controllers (thin, extract req → call one service → res.json({data})) → services (all business rules, RBAC, transactions, AI prompts) → repositories/clients (Supabase/Firebase/Gemini/Storage). Middlewares (auth, rbac, validate, rateLimit, requestContext, error) sit before routes; utils (ids, pagination, logger, AppError) are cross-cutting."),
        ("What is the frontend architecture?",
         "App.jsx does lazy code-splitting (Suspense + DeepamLoader) and guard nesting RequireAuth → RequireOnboarded → RequireRole. Contexts hold auth+language; services (typed apiClient modules) are the only place that talks to /api; Query owns caching; components are dumb — no business logic outside services."),
        ("Why backend-first API design?",
         "Every screen is an API contract first — web, future mobile/PWA, and AI agents all reuse the same REST. Lets us build frontend and backend in parallel and swap UI without touching rules."),
        ("How does the SPA routing work with Vercel?",
         "Vercel serves dist/ static. vercel.json rewrites /(.*) → /index.html so React Router can resolve /dashboard etc. on hard refresh. Without it, direct deep links would 404."),
        ("What is the request lifecycle for an authenticated read (e.g., GET /students/me)?",
         "Browser: apiClient interceptor does auth.currentUser.getIdToken() → Authorization: Bearer … → Express: helmet→cors→json→requestContext (requestId)→generalLimiter→route middleware: authenticate (verifyIdToken → upsert stub user if new → disabled check → attach req.user), requireRole(student), requireProfileComplete, controller getMe → profileService.getStudentProfile → studentsRepo.getById → res.json({data, meta:{requestId}}) → React Query caches 5m."),
        ("Why subpath imports like #services/*?",
         "Defined in backend/package.json imports — avoids ../../ hell, makes renames trivial, and enforces layering (e.g., controllers can’t deep-import repositories without it being obvious)."),
    ]),
    ("C. Database & Storage (6 Qs)", [
        ("Why Supabase PostgreSQL and not Firestore?",
         "Docs are legacy Firestore — runtime is PostgreSQL for: strong relations (FKs, UNIQUE student+job), GIN skill search, composite indexes, transactions, triggers, free-tier pg, and cheaper complex aggregations (analytics summary in one parallel count). Firestore pricing/complex queries would have hurt."),
        ("List the main tables and a key constraint.",
         "users (uid PK = Firebase UID), students/employers (1:1 FK), jobs, applications (PK application_id = `${student}_${job}` + UNIQUE(student_id,job_id) — one active app per job), notifications, analytics_events, platform_config singleton, chat_threads/messages, career_roadmaps/roadmap_progress, candidate_matches (job,student unique). All have updated_at triggers."),
        ("How do you handle skill search?",
         "students.searchable_skills TEXT[] GIN + jobs.required_skills GIN. Query via .overlaps() or .contains() — no LIKE scan. Used in candidateMatching and computeJobMatch."),
        ("How is file upload done without proxying through Node?",
         "POST /uploads/sign validates kind/mime/size (Zod+UPLOAD_POLICY) → Supabase createSignedUploadUrl (15 min) → frontend PUTs directly to Supabase bucket → POST /uploads/commit stores URL/path in DB and, for resumes, pdf-parse extracts keywords/text. Large files never touch Express memory beyond Multer for resume-analyzer PDF."),
        ("What are the storage buckets?",
         "resumes (private, 5 MB PDF, certs too), profile-pictures (public, 2 MB PNG/JPEG/WEBP), company-logos (public, 2 MB). Paths namespaced by uid: resumes/{uid}/resume-{uuid}.pdf etc."),
        ("How do you keep data consistent on apply?",
         "applications id is deterministic (student_job) so duplicate inserts conflict → 409 CONFLICT. Service creates application + notification (employer) in sequence and fires candidateMatching async; statusHistory is appended atomically. For admin verify, same pattern — verified_by_admin + verified_at together."),
    ]),
    ("D. Auth & Security (5 Qs)", [
        ("How does Firebase Auth integrate with your own DB?",
         "Firebase is identity only. Frontend mints ID token (getIdToken). Backend Firebase Admin verifyIdToken → decoded uid/email/picture. Then SELECT users by uid; if missing, upsert stub (role null, profileComplete false). If disabled → 403. req.user is then the DB user, not just Firebase claims. Role/profile lives in PostgreSQL."),
        ("Explain the 3 roles and how RBAC works.",
         "Roles: student, employer, admin (and null before onboarding). Middlewares requireRole(...allowed) and requireProfileComplete. Every service re-checks ownership (e.g., employer can list only own jobs, mutate only own jobId, updateStatus only if job belongs to employer). Admin guard checks role admin only."),
        ("How is admin created? Is there a separate table?",
         "No separate table. Promoted via POST /auth/admin-login: verify token, check email.toLowerCase() === env ADMIN_EMAIL (e.g., admin@lanturn.in), then users.update({role:'admin', profileComplete:true}). One env secret controls it."),
        ("How do you handle rate-limiting and validation?",
         "express-rate-limit 100/15s general, auth burst stricter, AI 20/day tracked in ai_usage (429 RATE_LIMITED). Every route has Zod schema — validates body/query/params, strips unknowns, returns VALIDATION_ERROR 400 with details[]."),
        ("How do you prevent IDOR/prompt injection/mass assignment?",
         "IDOR: services check resource.owner == req.user.uid before read/write. Mass assignment: USER_LOCKED_FIELDS [role,status,email,uid] ignored; Zod strict. Prompt injection: System-first prompts, delimiters around user text, temperature low (0.1 for scoring), responseMimeType application/json, explicit trailing ignore-instructions line."),
    ]),
    ("E. Features Deep-Dive (6 Qs)", [
        ("Walk through a student applying to a job.",
         "Internal Jobs: search/filter → cards sorted by computeJobMatch (skillMatchRatio*75 + bonuses for resume/projects/exp) → dialog shows Qualification Match (matched/missing chips) + Company Profile toggle → Apply Now → POST /jobs/:jobId/applications {coverLetter} → service checks no duplicate active app, creates row (snapshots name/photo/resume/skills), inserts notification (employer application_received), invalidates analytics, fires candidateMatchingService.generateAndStoreMatchScore async (builds jobData+candidateData → Gemini temp 0.1 → upsert candidate_matches). Student sees toast + button becomes Applied; Kanban gets new card."),
        ("What does computeJobMatch do?",
         "Pure frontend function in JobsPage — collects studentSkillSet from searchableSkills/professional.skills/resumeKeywords (lowercased). If job has requiredSkills, compute matched/missing via exact or substring includes, skillMatchRatio*75 + 5 each for resumeUrl/projects/experience/>5 skills/matched>0 clamped 15–100 → labels Highly Qualified ≥80 etc. Else text match on title+description: 50+matches*8."),
        ("Explain Career Aisle (57 roadmaps).",
         "Static CAREER_DOMAINS in backend/src/data/careerRoadmaps.js (frontend/mobile/AI/cloud/product… 57). User enrolls (POST /roadmaps/me/enroll) → career_roadmaps row + roadmap_progress rows. Timeline renders alternating StagePanels with topic checkboxes → PATCH /roadmaps/me/:id/progress toggles. Resume Analyzer uploads PDF → resumeAnalyzer (Gemini) returns matchedTopicKeys → Sync bulk completes those topics + recommends start stage."),
        ("What are Candidate Matches and Company/Career DNA?",
         "CandidateMatches: On apply, Gemini scores 0–100 weighted Technical 40/Projects 25/Experience 20/Education 15 + missingSkills/recommendations → stored per job×student → employer sees on applicants page + dashboard recommendations. Career DNA: Student’s 6-dimension radar (Gemini tailored labels) persisted in students.professional.careerDna, regenerated via /ai/career-dna. Company DNA: Employer’s 6-dimension personality built from company fields + 10 active jobs, cooldown 24h, visible to students in job dialog."),
        ("How does Kanban work for employer applicants?",
         "KanbanBoard is generic @dnd-kit component — columns = STATUS_CONFIG (Applied blue→Reviewing amber→Shortlisted emerald→Hired violet→Rejected red). Items are applications mapped to id=applicationId. getItemColumn returns status. onDragEnd calls statusMutation with toColumn → PATCH /applications/:id/status → toast + invalidate + notification to student (application_status). Also supports List view with same actions."),
        ("What does the Admin Portal allow?",
         "Dashboard (8 stat cards + Recent Posts/Users via analytics summary), Post Verification (Search+status filter, Verify verified_by_admin, Change Status moderateJob), Users (search+role/status filter, photo/initial, Active/Disabled pill, Actions dropdown: Change Role / Disable/Enable). Plus Platform Config and Analytics series (by type/days) — all behind requireRole(admin)."),
    ]),
    ("F. AI & Gemini (5 Qs)", [
        ("Which AI features use Gemini and how?",
         "All: Resume Review (score 0-100, ATS, strengths/weaknesses/fixes, keywordsMissing, predictedRole), Resume Match (matchScore, matched/missing, experienceFit), Skill Gap (importance+suggestion), Skill Test (generate 3 open-ended Qs tiered by rating + evaluate → medal gold/silver/bronze/basic/none), Cover Letter (tone), Career Chat (threaded, 10-history, <150w bullets, mode career_guidance), Career DNA (6 tailored dims), Resume Extract (personal/academic/professional/social), Job Description AI (HR copywriter), Company DNA, Employer Hiring Assistant (context jobs+pool), Candidate Matching."),
        ("How do you handle quota, retries, and safety for Gemini?",
         "GEMINI_API_KEY can be comma-separated — 429 rotates to next key; 5xx/timeout rotates to next model (flash-lite → flash → 2.5). temperature 0.1 (scoring) to 0.8 (chat), maxOutputTokens 2048, responseMimeType application/json for structured, safety BLOCK_ONLY_HIGH, Zod validates Gemini JSON → 502 UPSTREAM_ERROR on bad shape, ai_usage daily cap 20 → 429, 25s timeout with retry only on transient."),
        ("How is resume text handled?",
         "On upload/commit, pdf-parse extracts resumeText + keywords → cached in students.resume_text/resume_keywords. All AI calls then do fetchResumeText: if resume_url exists, getSignedDownloadUrl (1h) → fetch → pdf-parse → cache; else fallback to profile fields. If no PDF, 422."),
        ("What is the difference between Career DNA and Company DNA?",
         "Career DNA is student-facing — 6 dimensions tailored to the student’s skills/projects, produced by Gemini, stored in students.professional.careerDna, shown as radar + insights (strengths/weaknesses/recommendations). Company DNA is employer-facing — 6 dimensions tailored to company + its active jobs, stored in employers.company_dna, 24h cooldown, public to students in job dialog. Same shape, mirrored prompts."),
        ("Why not use OpenAI directly?",
         "GEMINI_INTEGRATION.md mandates Gemini only per constraints; also free tier is more generous and free-tier JSON mode + safety categories fit our use case. DeepSeek is the documented fallback if Gemini is down."),
    ]),
    ("G. Deployment & Challenges (5 Qs)", [
        ("How do you deploy?",
         "Frontend → Vercel: git push master → auto-build (npm run build) → dist/ + vercel.json SPA rewrite, env VITE_* set in dashboard. Backend → Render/Oracle: env PORT/CORS/GEMINI/SUPABASE/FIREBASE set, node src/server.js, health /api/health. Supabase hosts DB+Storage. Domain via Vercel/Render env."),
        ("What was the hardest bug and how did you fix it?",
         "Pick one you lived: e.g., Post Verification 500 because jobs.status had no 'verified' value — added migration 008 (verified_by_admin + verified_at + status check) and updated rowToJob + verifyJob service; or Kanban DnD not updating because getItemColumn returned wrong status — fixed by mapping id=applicationId; or Resume Review 422 because resumeText empty — added fetchResumeText fallback + pdf-parse caching."),
        ("How do you handle env secrets?",
         "Backends .env is gitignored, copied from .env.example, Zod-validated with fail-fast (process.exit 1 on bad, logs fieldErrors), replaces \\n in private key. Frontend VITE_ vars are not secrets (bundled), only the API URL + Firebase web config. Service role key never reaches browser."),
        ("How did you keep costs at $0?",
         "Entire stack is free tier: Firebase free (Auth), Supabase free (PG+Storage), Gemini free (multi-key rotation), Vercel free, Render free, no paid email/search — Jooble free key, email skipped if not configured. No RLS bills because service_role bypasses. AGENTS notes: no dotenv dep (custom loader) to keep deps tiny."),
        ("What would you build next?",
         "From MOBILE_PLAN/TODO: PWA (manifest + Workbox vite-plugin-pwa, stale-while-revalidate for GET /jobs, FCM push via /api/devices/register), real-time chat (Supabase Realtime onSnapshot for notifications + Bell), video interviews, referrals, ATS export, RAG over resume+job corpus, premium analytics. Also backfill-match-scores script already exists for historical jobs."),
    ]),
]

for section_title, qas in qa_groups:
    add_heading_styled(section_title, level=2, size=10.5, color=RGBColor(0x1A,0x1A,0x1A))
    for q, a in qas:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.left_indent = Inches(0.05)
        # Q
        r = p.add_run("Q:  " + q)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = BRAND
        r.font.name = 'Calibri'
        # A
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        p2.paragraph_format.left_indent = Inches(0.15)
        # add yellow left border via paragraph shading? Use simple indent + marker
        r2 = p2.add_run("A:  " + a)
        r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor(0x33,0x33,0x33)
        r2.font.name = 'Calibri'

# ────────────────────────────────────────────────────────────────
# 12. TIPS & CLOSING
# ────────────────────────────────────────────────────────────────
add_heading_styled("12  Presentation Tips & Closing", level=1)

add_heading_styled("How to Present (10 minutes max)", level=2, size=11)
add_bullet_list([
    "Start with the problem (30s): manual placements, bias, no AI feedback → LanTURN’s answer (Turning Learning into Career, free-tier).",
    "Stack slide (60s): Show the one-line stack table — emphasise Monorepo + Supabase PG + Firebase Auth + Gemini (free-tier story = strong viva point).",
    "Architecture (60s): Draw the Bearer flow on board — Browser → Express (Helmet/CORS/Zod) → Firebase verify → Services → Supabase/Gemini. Mention layering rule (routes→controllers→services→repos).",
    "Demo (5–6 min): Follow the Demo Flow in §10 — don’t click everything; pick one student journey (apply + AI review) and one employer journey (post + Kanban). Narrate the matchScore / pipeline colours.",
    "Close (30s): Free-tier deployment (Vercel+Supabase+Firebase), what’s next (PWA + Realtime Bell + Video), and “all code is API-first so mobile reuses 100%”.",
], size=7.5)

add_heading_styled("Viva Defence Lines", level=3, size=9.5)
add_bullet_list([
    "If asked “Why not Firestore?” — answer with GIN search, UNIQUE constraint, composite indexes, cheaper aggregations, and that docs are legacy.",
    "If challenged on yellow background / monochrome UI — answer: Nothing OS inspired, brand-900 + accent #FFC107 + brutalist shadows (hard shadows + 2px border) — intentional, not default Tailwind.",
    "If asked about scalability — mention stateless Express, cursor pagination (limit+1 → nextCursor), parallel analytics counts, and that Supabase scales vertically before needing read replicas.",
    "If they test security — walk through disabled check, ownership re-check in services, Zod + locked fields, signed URLs, env-only secrets.",
], size=7.5)

add_heading_styled("One-Line Answers to Keep Ready", level=3, size=9.5)
one_liners = [
    "“What’s LanTURN?” — An AI-native, free-tier, API-first campus placement platform with 3 roles and 8 AI tools that turns resumes into roadmaps and jobs into matches.",
    "“Why should we hire it?” — It cuts screening time (auto matchScore + shortlist), gives students instant feedback (review/skill gap/career DNA), and gives admin full governance (verify/pause/role).",
    "“What did you learn?” — Building backend-first, validating with Zod, handling Gemini quota/retries, and keeping a monorepo shippable on free infra.",
]
add_bullet_list(one_liners, size=7.5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(2)
# yellow divider
r = p.add_run("─" * 48)
r.font.size = Pt(8)
r.font.color.rgb = ACCENT
r.font.name = 'Calibri'
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(1)
r = p2.add_run("All the best for tomorrow — you’ve built a real product. Open this doc and the app side-by-side, breathe, and narrate the journey.")
r.font.size = Pt(9)
r.bold = True
r.font.color.rgb = BRAND
r.font.name = 'Calibri'
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p3.add_run("LanTURN  •  Siva Sivani Degree College  •  " + datetime.date.today().strftime("%Y"))
r.font.size = Pt(7.5)
r.font.color.rgb = GREY
r.font.name = 'Calibri'

# ── Footers with page number ──
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('w:space'), 'preserve'); instr.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar1); r._r.append(instr); r._r.append(fldChar2)
    r2 = p.add_run("  •  LanTURN  •  Confidential — for presentation only")
    r2.font.size = Pt(7)
    r2.font.color.rgb = GREY
    r2.font.name = 'Calibri'

out = r"C:\Users\Aravind\Downloads\clg projects\New folder\LanTURN_Presentation_Guide.docx"
doc.save(out)
print(f"Saved to {out}")
